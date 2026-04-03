"""Generate structured DOCX briefing from JSON data."""
import json
import logging
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

# Color scheme
COLOR_HEADER = RGBColor(0x1B, 0x3A, 0x5C)  # dark navy
COLOR_RISK_HIGH = RGBColor(0xCC, 0x00, 0x00)
COLOR_RISK_MED = RGBColor(0xCC, 0x88, 0x00)
COLOR_RISK_LOW = RGBColor(0x00, 0x88, 0x00)
COLOR_GREY = RGBColor(0x66, 0x66, 0x66)


def generate_briefing_docx(json_data: dict, model_name: str, output_path: Path, filenames: list[str]):
    """Generate a structured DOCX from parsed tender JSON."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === TITLE ===
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BRIEFING PRZETARGOWY")
    run.font.color.rgb = COLOR_HEADER
    run.font.size = Pt(18)

    # Tender name subtitle
    nazwa = json_data.get("nazwa", "Brak nazwy")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(nazwa)
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_HEADER
    run.bold = True

    # Metadata line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Model: {model_name}")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GREY

    doc.add_paragraph()  # spacer

    # === 1. INFORMACJE PODSTAWOWE ===
    _add_section_heading(doc, "1. INFORMACJE PODSTAWOWE")

    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_info_row(info_table, "Zamawiający", json_data.get("zamawiajacy"))
    _add_info_row(info_table, "Nr postępowania", json_data.get("numer_postepowania"))
    _add_info_row(info_table, "Link", json_data.get("link"))

    doc.add_paragraph()

    # === 2. TERMINY ===
    _add_section_heading(doc, "2. TERMINY")

    terminy = json_data.get("terminy", {})
    t_table = doc.add_table(rows=0, cols=2)
    t_table.style = "Light Grid Accent 1"
    t_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_info_row(t_table, "Składanie ofert", terminy.get("skladanie_ofert"))
    _add_info_row(t_table, "Realizacja", terminy.get("realizacja"))
    _add_info_row(t_table, "Gwarancja", terminy.get("gwarancja"))
    _add_info_row(t_table, "Związanie ofertą", terminy.get("zwiazanie_oferta"))

    doc.add_paragraph()

    # === 3. FINANSE ===
    _add_section_heading(doc, "3. FINANSE")

    finanse = json_data.get("finanse", {})
    f_table = doc.add_table(rows=0, cols=2)
    f_table.style = "Light Grid Accent 1"
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_info_row(f_table, "Wadium", finanse.get("wadium"))
    _add_info_row(f_table, "Zabezpieczenie NWU", finanse.get("zabezpieczenie_nwu"))
    _add_info_row(f_table, "Szacunkowa wartość", finanse.get("szacunkowa_wartosc"))

    doc.add_paragraph()

    # === 4. TRYB ===
    _add_section_heading(doc, "4. TRYB POSTĘPOWANIA")

    tryb = json_data.get("tryb", {})
    tr_table = doc.add_table(rows=0, cols=2)
    tr_table.style = "Light Grid Accent 1"
    tr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_info_row(tr_table, "Otwarcie", tryb.get("otwarcie"))
    _add_info_row(tr_table, "Tryb prawny", tryb.get("tryb_prawny"))
    _add_info_row(tr_table, "Aukcja/Negocjacje", tryb.get("aukcja_negocjacje"))
    _add_info_row(tr_table, "Konsorcjum", tryb.get("konsorcjum"))

    doc.add_paragraph()

    # === 5. WYMAGANIA - DOŚWIADCZENIE ===
    _add_section_heading(doc, "5. WYMAGANIA — DOŚWIADCZENIE")
    _add_bullet_list(doc, json_data.get("wymagania_doswiadczenie", []))

    # === 6. WYMAGANIA - KADRA ===
    _add_section_heading(doc, "6. WYMAGANIA — KADRA")
    _add_bullet_list(doc, json_data.get("wymagania_kadra", []))

    # === 6b. WYMAGANIA — DODATKOWE ===
    wym_dod = json_data.get("wymagania_dodatkowe", [])
    if wym_dod:
        _add_section_heading(doc, "6b. WYMAGANIA — DODATKOWE")
        _add_bullet_list(doc, wym_dod)

    # === 7. ZAKRES PRAC ===
    _add_section_heading(doc, "7. ZAKRES PRAC")
    zakres = json_data.get("zakres_prac", "Brak danych")
    doc.add_paragraph(str(zakres))

    # Quick scope summary from stats (key numbers at a glance)
    stats = json_data.get("podsumowanie_statystyczne", {})
    if stats:
        quick_items = []
        for key, label in [("slupy", "Słupy"), ("fundamenty", "Fundamenty"),
                           ("przewody_fazowe_m", "Przewody fazowe"), ("opgw_m", "OPGW"),
                           ("izolatory", "Izolatory")]:
            val = stats.get(key)
            if val:
                quick_items.append(f"{label}: {val}")
        if quick_items:
            p = doc.add_paragraph()
            run = p.add_run("Kluczowe ilości: ")
            run.bold = True
            run.font.size = Pt(9)
            run2 = p.add_run(" | ".join(quick_items))
            run2.font.size = Pt(9)
            run2.font.color.rgb = COLOR_GREY

    # === 7a. LOKALIZACJA ===
    lokalizacja = json_data.get("lokalizacja", {})
    if lokalizacja and isinstance(lokalizacja, dict):
        _add_section_heading(doc, "7a. LOKALIZACJA")

        lok_table = doc.add_table(rows=0, cols=2)
        lok_table.style = "Light Grid Accent 1"
        lok_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        gminy = lokalizacja.get("gminy", [])
        if gminy and isinstance(gminy, list):
            _add_info_row(lok_table, "Gminy", ", ".join(str(g) for g in gminy))
        powiat = lokalizacja.get("powiat")
        if powiat:
            _add_info_row(lok_table, "Powiat", powiat)
        woj = lokalizacja.get("wojewodztwo")
        if woj:
            _add_info_row(lok_table, "Województwo", woj)
        punkty = lokalizacja.get("punkty_charakterystyczne", [])
        if punkty and isinstance(punkty, list):
            _add_info_row(lok_table, "Punkty charakterystyczne", ", ".join(str(p) for p in punkty))
        teren = lokalizacja.get("opis_terenu")
        if teren:
            _add_info_row(lok_table, "Opis terenu", teren)

        # Google Maps link
        maps_query = _build_maps_query(gminy, punkty, json_data.get("nazwa", ""))
        if maps_query:
            maps_url = f"https://www.google.com/maps/search/{quote(maps_query)}"
            _add_info_row(lok_table, "Google Maps", maps_url)

        doc.add_paragraph()

    # === 7b. PROJEKT TECHNICZNY ===
    projekt = json_data.get("projekt_techniczny", {})
    if projekt and isinstance(projekt, dict) and any(v for v in projekt.values() if v):
        _add_section_heading(doc, "7b. PROJEKT TECHNICZNY")

        pr_table = doc.add_table(rows=0, cols=2)
        pr_table.style = "Light Grid Accent 1"
        pr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        _add_info_row(pr_table, "Projekt istnieje", projekt.get("czy_istnieje"))
        firma = projekt.get("firma_projektowa")
        if firma:
            _add_info_row(pr_table, "Firma projektowa", firma)
        zakres_pr = projekt.get("zakres_projektu")
        if zakres_pr:
            _add_info_row(pr_table, "Zakres projektu", zakres_pr)
        uwagi_pr = projekt.get("uwagi")
        if uwagi_pr:
            _add_info_row(pr_table, "Uwagi", uwagi_pr)

        doc.add_paragraph()

    # === 7c. OBOWIĄZKI WYKONAWCY ===
    obowiazki = json_data.get("obowiazki_wykonawcy", [])
    if obowiazki:
        _add_section_heading(doc, "7c. OBOWIĄZKI WYKONAWCY (poza standardowymi robotami)")
        _add_bullet_list(doc, obowiazki)

    # === 8. ZESTAWIENIE MATERIAŁÓW ===
    _add_section_heading(doc, "8. ZESTAWIENIE MATERIAŁÓW")

    materialy = json_data.get("zestawienie_materialow", [])
    if materialy:
        for mat in materialy:
            if not isinstance(mat, dict):
                continue

            # Category sub-heading
            kategoria = str(mat.get("kategoria", ""))
            lacznie = str(mat.get("lacznie_sztuk", ""))
            count = str(mat.get("pozycje_count", ""))

            p = doc.add_paragraph()
            run = p.add_run(f"{kategoria}")
            run.bold = True
            run.font.size = Pt(11)
            if lacznie:
                run2 = p.add_run(f"  —  łącznie: {lacznie}")
                run2.font.color.rgb = COLOR_GREY
                run2.font.size = Pt(9)
            if count:
                run3 = p.add_run(f"  ({count} pozycji)")
                run3.font.color.rgb = COLOR_GREY
                run3.font.size = Pt(9)

            # List all items
            kluczowe = mat.get("kluczowe_pozycje", [])
            if isinstance(kluczowe, list):
                for item in kluczowe:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif kluczowe:
                doc.add_paragraph(str(kluczowe), style="List Bullet")
    else:
        doc.add_paragraph("Brak danych o materiałach.")

    doc.add_paragraph()

    # === 9. PODSUMOWANIE STATYSTYCZNE ===
    stats = json_data.get("podsumowanie_statystyczne", {})
    if stats and any(v for v in stats.values() if v):
        _add_section_heading(doc, "9. PODSUMOWANIE STATYSTYCZNE")
        s_table = doc.add_table(rows=0, cols=2)
        s_table.style = "Light Grid Accent 1"
        s_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        stat_labels = {
            "slupy": "Słupy",
            "fundamenty": "Fundamenty",
            "izolatory": "Izolatory",
            "przewody_fazowe_m": "Przewody fazowe",
            "opgw_m": "OPGW",
            "tlumiki_drgan": "Tłumiki drgań",
            "zawiesia": "Zawiesia",
            "uziemienia": "Uziemienia",
            "oznakowanie": "Oznakowanie",
            "inne": "Inne",
        }
        for key, label in stat_labels.items():
            val = stats.get(key)
            if val:
                _add_info_row(s_table, label, val)

        doc.add_paragraph()

    # === 10. KRYTERIA OCENY ===
    kryteria = json_data.get("kryteria_oceny", {})
    if kryteria and any(v for v in kryteria.values() if v):
        _add_section_heading(doc, "10. KRYTERIA OCENY OFERT")
        k_table = doc.add_table(rows=0, cols=2)
        k_table.style = "Light Grid Accent 1"
        k_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        _add_info_row(k_table, "Waga ceny", kryteria.get("cena_waga"))
        inne_kry = kryteria.get("inne_kryteria", [])
        if inne_kry and isinstance(inne_kry, list):
            _add_info_row(k_table, "Inne kryteria", "; ".join(str(k) for k in inne_kry))
        opis_kry = kryteria.get("opis")
        if opis_kry:
            _add_info_row(k_table, "Opis", opis_kry)

        doc.add_paragraph()

    # === 11. KARY UMOWNE ===
    kary = json_data.get("kary_umowne", [])
    if kary:
        _add_section_heading(doc, "11. KARY UMOWNE")
        _add_bullet_list(doc, kary)

    # === 12. ZMIANY Z ODPOWIEDZI NA PYTANIA ===
    zmiany = json_data.get("zmiany_z_odpowiedzi", [])
    if zmiany:
        _add_section_heading(doc, "12. ZMIANY Z ODPOWIEDZI NA PYTANIA")
        _add_bullet_list(doc, zmiany)

    # === 13. RYZYKA I UWAGI ===
    _add_section_heading(doc, "13. RYZYKA I UWAGI")

    ryzyka = json_data.get("ryzyka_i_uwagi", [])
    if ryzyka:
        for ryzyko in ryzyka:
            if not isinstance(ryzyko, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(ryzyko))
                continue

            poziom = str(ryzyko.get("poziom", "")).upper()
            opis = ryzyko.get("opis", "")
            zrodlo = ryzyko.get("zrodlo", "")

            p = doc.add_paragraph(style="List Bullet")

            # Color-coded level tag
            tag_run = p.add_run(f"[{poziom}] ")
            tag_run.bold = True
            if "WYSOKI" in poziom or "HIGH" in poziom:
                tag_run.font.color.rgb = COLOR_RISK_HIGH
            elif "ŚREDNI" in poziom or "MEDIUM" in poziom or "MED" in poziom:
                tag_run.font.color.rgb = COLOR_RISK_MED
            else:
                tag_run.font.color.rgb = COLOR_RISK_LOW

            p.add_run(str(opis))

            if zrodlo:
                src_run = p.add_run(f" [{zrodlo}]")
                src_run.font.color.rgb = COLOR_GREY
                src_run.font.size = Pt(8)
    else:
        doc.add_paragraph("Nie zidentyfikowano szczególnych ryzyk.")

    # === 14. REKOMENDACJE ===
    _add_section_heading(doc, "14. REKOMENDACJE")
    rekomendacje = json_data.get("rekomendacje", [])
    if rekomendacje:
        for i, rek in enumerate(rekomendacje, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True
            p.add_run(str(rek))
    else:
        doc.add_paragraph("Brak rekomendacji.")

    # === FOOTER ===
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("─" * 40)
    run.font.color.rgb = COLOR_GREY
    run.font.size = Pt(8)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(
        f"Źródła: {', '.join(filenames)}\n"
        f"Model: {model_name} | Wygenerowano automatycznie"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = COLOR_GREY

    # Save
    doc.save(str(output_path))
    logger.info(f"DOCX saved: {output_path}")


def _add_section_heading(doc, text: str):
    """Add a styled section heading."""
    heading = doc.add_heading(level=2)
    run = heading.add_run(text)
    run.font.color.rgb = COLOR_HEADER
    run.font.size = Pt(13)


def _add_info_row(table, label: str, value):
    """Add a label-value row to a 2-column table."""
    if value is None:
        value = "—"
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    # Bold the label
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True


def _add_bullet_list(doc, items: list):
    """Add a bulleted list."""
    if not items:
        doc.add_paragraph("Brak danych.")
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _build_maps_query(gminy: list, punkty: list, nazwa: str) -> str:
    """Build a Google Maps search query from location data."""
    parts = []
    if punkty:
        parts.extend(str(p) for p in punkty[:3])
    elif gminy:
        parts.extend(str(g) for g in gminy[:3])
    if not parts and nazwa:
        parts.append(nazwa)
    return " ".join(parts) if parts else ""


def parse_llm_response(raw_response: str) -> dict:
    """Parse JSON from LLM response, handling markdown code blocks."""
    text = raw_response.strip()

    # Strip markdown code block if present
    if text.startswith("```"):
        # Remove first line (```json or ```)
        lines = text.split("\n")
        text = "\n".join(lines[1:])
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse JSON: {e}")
        logger.debug(f"Raw response (first 500 chars): {text[:500]}")

        # Try to find JSON object in the text
        start = text.find("{")
        end = text.rfind("}") + 1
        if start >= 0 and end > start:
            try:
                return json.loads(text[start:end])
            except json.JSONDecodeError:
                pass

        raise ValueError(f"Could not parse LLM response as JSON: {e}")
